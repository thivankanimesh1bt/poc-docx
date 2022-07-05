import json
import docx
from pptx.util import Pt
import time

import boto3
from botocore.exceptions import ClientError
import logging
from io import BytesIO

import asyncio
from docxcompose.composer import Composer

async def process(doc, data):
    doc.add_heading('gayashan', 0)

    i = 0
    while i < 1500:
        table = doc.add_table(rows=1, cols=6, style='Colorful List')
        row = table.rows[0].cells
        row[0].text = 'Id'
        row[1].text = 'Name'
        row[2].text = 'aaa'
        row[3].text = 'bbbb'
        row[4].text = 'cccc'
        row[5].text = 'dddd'


        for id, name in data:
            row = table.add_row().cells
            row[0].text = str(id)
            row[1].text = "name"
            row[2].text = "name 1"
            row[3].text = "nam  2"
            row[3].text = "name 3"
            row[4].text = name


        paragraph = doc.add_paragraph(' ')
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(5)
        i += 1

    return doc

async def run():
    doc1 = docx.Document()
    doc2 = docx.Document()
    doc3 = docx.Document()
    doc4 = docx.Document()

    start_time = time.perf_counter ()
    
    data = (
        (1, 'gaya 1'),
        (2, 'gaya 2'),
        (3, 'gaya 3'),
        (1, 'gaya 4'),
        (2, 'gaya 5'),
        (3, 'gaya 6'),
        (1, 'gaya 7'),
        (2, 'gaya 8'),
        (3, 'gaya 9'),
        (1, 'gaya 10'),
        (1, 'gaya 11'),
        (2, 'gaya 12'),
        (3, 'gaya 13'),
        (1, 'gaya 14'),
        (2, 'gaya 15'),
        (3, 'gaya 16'),
        (1, 'gaya 17'),
        (2, 'gaya 18'),
        (3, 'gaya 19'),
        (1, 'gaya 20')
    )

    # newDocs = await asyncio.gather(process(doc1, data), process(doc2, data), process(doc3, data), process(doc4, data)) 
    newDocs = await asyncio.gather(process(doc1, data)) 


    # combine docx
    combine = docx.Document()
    composer = Composer(combine)
    composer.append(newDocs[0])
    # composer.append(newDocs[1])
    # composer.append(newDocs[2])
    # composer.append(newDocs[3])

    # upload to s3
    s3_client = boto3.client('s3')

    try:
        with BytesIO() as fileobj:
            composer.save(fileobj)
            fileobj.seek(0)
            res = s3_client.upload_fileobj(fileobj, 'poc-docx', 'output.docx')
    except ClientError as e:
        logging.error(e)
        return False


def trigger(event, context):
    for record in event['Records']:
        payload = record["body"]
        print(payload)
        asyncio.run(run())
        print("done")


def docxGenerate(event, context):
    sqs_client = boto3.client("sqs", region_name="us-east-1")

    message = {"key": "value"}

    response = sqs_client.send_message(
        QueueUrl="https://sqs.us-east-1.amazonaws.com/954784122482/poc-queue.fifo",
        MessageBody=json.dumps(message),
        MessageGroupId='docx-generate'
    )
    
    response = {
        "statusCode": 200,
        "body": "Success"
    }

    return response