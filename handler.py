import json

from docx import Document
from pptx.util import Inches
from pptx.util import Pt
from pptx.dml.color import RGBColor
import copy

import boto3
from botocore.exceptions import ClientError
import logging
from io import BytesIO

def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)

def replace_txt(replacements, document):
    for paragraph in document.paragraphs:
        for match, replacement in replacements.items():
            replace_text_in_paragraph(paragraph, match, replacement)

    for section in document.sections:
        footer = section.footer
        header = section.header
        for paragraph in footer.paragraphs:
            for match, replacement in replacements.items():
                replace_text_in_paragraph(paragraph, match, replacement)

        for paragraph in header.paragraphs:
            for match, replacement in replacements.items():
                replace_text_in_paragraph(paragraph, match, replacement)

    #how to replace tags inside table
    # for table in document.tables:
    #     for row in table.rows:
    #         for cell in row.cells:
    #             for paragraph in cell.paragraphs:
    #                 for match, replacement in replacements.items():
    #                     replace_text_in_paragraph(paragraph, match, replacement)


def add_data_to_table(table):
    data = (
        (1, 'gaya 1'),
        (2, 'gaya 2'),
        (3, 'gaya 3'),
        (1, 'gaya 4'),
        (2, 'gaya 5'),
        (3, 'gaya 6'),
        (1, 'gaya 7'),
        (2, 'gaya 8'),
        (3, 'gaya 9'),
        (1, 'gaya 10'),
        (1, 'gaya 11'),
        (2, 'gaya 12'),
        (3, 'gaya 13'),
        (1, 'gaya 14'),
        (2, 'gaya 15'),
        (3, 'gaya 16'),
        (1, 'gaya 17'),
        (2, 'gaya 18'),
        (3, 'gaya 19'),
        (1, 'gaya 20')
    )

    print(table.rows[0].cells)

    row = table.rows[0].cells
    row[0].text = 'Id'
    row[1].text = 'Name'
    row[2].text = 'aaa'
    row[3].text = 'bbbb'
    row[4].text = 'cccc'

    for id, name in data:
        row = table.add_row().cells
        row[0].text = str(id)
        row[1].text = "name"
        row[2].text = "name 1"
        row[3].text = "nam  2"
        row[4].text = "name 3"

    print("table style===========>",table.style)

def update_table(replacements, document):  
    print("template table style ===========>",document.tables[0].style)
    #add data to existing table
    add_data_to_table(document.tables[0])

    new_tbl = copy.deepcopy(document.tables[0])
    print("copied table============>",new_tbl)

    #table = document.add_table(rows=1, cols=6, style=)
    paragraph = document.add_paragraph(' ')
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(5)

    #create new table using existing styles
    table = document.add_table(rows=1, cols=5)
    table.style = document.tables[0].style
    add_data_to_table(table)
    paragraph = document.add_paragraph(' ')
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(5)


    table_2 = document.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    add_data_to_table(table_2)


def docxGenerate(event, context):
    doc = Document('input.docx')
    dataDump = {
        "text_replaces":  {
            '{{var1}}': "Example Name",
            '{{var2}}': 'External asset workflow',
            '{{var3}}': "302929393",
            '{{var4}}': 'LGIM/Client user',
            '{{var5}}': "+972-5056000000",
            '{{var6}}': "example@example.com",
            '{{var7}}': 'Active Client page ??? External assets',
            '{{var8}}': 'Information button text',
            '{{var10}}': "03 Jan, 2022",
            '{{var11}}': "10,000",
            '{{var12}}': "3,000",
            '{{var13}}':"7,000",
            },
        "table_replaces": {
            "cashFlows":{
                "headers": ["cashflow year","cashflow fixed","cashflow real"],
                "rows": [ 
                    {
                        "cashflow_year": "2020","cashflow_fixed":"1","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"2","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"3","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"4","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"5","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"6","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"7","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"8","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"9","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"10","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"11","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"12","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"13","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"14","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"15","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"16","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"17","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"18","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"19","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"20","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"21","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"22","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"23","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"24","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"25","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"26","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"27","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"28","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"29","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"30","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"31","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"32","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"33","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"34","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"35","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"36","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"37","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"38","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"39","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"40","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"41","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"42","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"43","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"44","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"45","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"51","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"52","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"53","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"54","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"55","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"56","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"57","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"58","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"59","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"510","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"151","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"512","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"153","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"154","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"155","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"61","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"62","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"63","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"64","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"65","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"66","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"67","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"68","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"69","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"170","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"171","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"172","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"173","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"174","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"175","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"17","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"82","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"83","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"84","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"85","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"86","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"87","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"88","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"89","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"190","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"191","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"192","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"193","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"194","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"195","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"17","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"82","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"83","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"84","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"85","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"86","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"87","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"88","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"89","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"190","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"191","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"192","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"193","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"194","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"195","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"1","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"2","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"3","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"4","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"5","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"6","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"7","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"8","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"9","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"10","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"11","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"12","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"13","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"14","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"15","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"16","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"17","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"18","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"19","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"20","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"21","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"22","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"23","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"24","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"25","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"26","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"27","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"28","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"29","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"30","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"31","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"32","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"33","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"34","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"35","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"36","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"37","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"38","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"39","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"40","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"41","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"42","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"43","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"44","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"45","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"51","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"5112","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"5113","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"5114","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"5115","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"5116","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"5117","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"58","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"5911","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"51110","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"15111","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"5112","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"1513","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"1514","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"1515","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"611","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"612","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"6113","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"614","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"615","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"616","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"617","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"618","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"619","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"1170","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"1171","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"1172","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"1173","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"1714","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"1715","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"1117","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"812","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"813","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"814","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"815","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"816","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"817","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"818","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"819","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"1910","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"1911","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"1912","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"1913","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"1914","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"1915","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"117","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"812","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"813","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"814","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"815","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"816","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"187","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"818","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"189","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"1190","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"1911","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"1192","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"1913","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"1941","cashflow_real": "123123"
                    },
                    {
                        "cashflow_year": "2020","cashflow_fixed":"1915","cashflow_real": "123123"
                    }
                 ]
            }
        } 
             
    }

    response = {
        "statusCode": 200,
        "body": "Success"
    }

    replace_txt(dataDump["text_replaces"], doc)
    update_table(dataDump["table_replaces"]["cashFlows"], doc)
    # doc.save('output.docx')

    # upload to s3
    s3_client = boto3.client('s3')

    try:
        with BytesIO() as fileobj:
            doc.save(fileobj)
            fileobj.seek(0)
            res = s3_client.upload_fileobj(fileobj, 'poc-docx', 'output.docx')
    except ClientError as e:
        logging.error(e)
        return False

    return response


